"""Belge okuma modülü.

PDF, DOCX, TXT, MD, HTML dosyalarından ve URL'lerden metin çıkarır.
PDF metni çıkarılırken sayfa numaraları `--- PAGE X ---` işaretleriyle korunur.
"""

from __future__ import annotations

from pathlib import Path
from urllib.parse import urljoin, urlparse

import requests

from .config import SUPPORTED_EXTENSIONS
from .utils import clean_text, cleanup_temp_file, logger, save_temp_file

_REQUEST_TIMEOUT = 30
_USER_AGENT = (
    "Mozilla/5.0 (compatible; MechanicalSpecAnalyzer/1.0; +https://example.local)"
)
_MAX_DOWNLOAD_BYTES = 60 * 1024 * 1024  # 60 MB güvenlik sınırı


class DocumentReadError(Exception):
    """Belge okuma sırasında oluşan kullanıcı dostu hata."""


# --------------------------------------------------------------------------- #
# Format bazlı çıkarıcılar
# --------------------------------------------------------------------------- #
def extract_text_from_pdf(path: Path) -> str:
    """PyMuPDF (fitz) ile PDF'ten metin çıkarır. Sayfa işaretleri eklenir."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover
        raise DocumentReadError(
            "PyMuPDF (pymupdf) kurulu değil. 'pip install pymupdf' çalıştırın."
        ) from exc

    parts: list[str] = []
    try:
        with fitz.open(path) as doc:
            for index, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                parts.append(f"--- PAGE {index} ---\n{page_text}")
    except Exception as exc:
        raise DocumentReadError(f"PDF okunamadı: {exc}") from exc

    text = "\n".join(parts)
    if not text.strip():
        raise DocumentReadError(
            "PDF'ten metin çıkarılamadı. Belge taranmış (görüntü) olabilir; "
            "OCR'lı bir PDF deneyin."
        )
    return text


def extract_text_from_docx(path: Path) -> str:
    """python-docx ile DOCX'ten metin çıkarır (paragraflar + tablolar)."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise DocumentReadError(
            "python-docx kurulu değil. 'pip install python-docx' çalıştırın."
        ) from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocumentReadError(f"DOCX okunamadı: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    # Tablolardaki metni de topla
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise DocumentReadError("DOCX dosyasından metin çıkarılamadı (boş belge).")
    return text


def extract_text_from_txt(path: Path) -> str:
    """TXT/MD dosyasını UTF-8 okur, hata olursa errors='ignore' kullanır."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        raise DocumentReadError(f"Metin dosyası okunamadı: {exc}") from exc

    if not text.strip():
        raise DocumentReadError("Metin dosyası boş.")
    return text


def extract_text_from_html(path: Path) -> str:
    """BeautifulSoup ile HTML'den script/style temizleyerek metin çıkarır."""
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:  # pragma: no cover
        raise DocumentReadError(
            "beautifulsoup4 kurulu değil. 'pip install beautifulsoup4' çalıştırın."
        ) from exc

    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        raise DocumentReadError(f"HTML dosyası okunamadı: {exc}") from exc

    return _html_to_text(raw)


def _html_to_text(raw_html: str) -> str:
    """Ham HTML string'ini temiz metne çevirir."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    if not text.strip():
        raise DocumentReadError("HTML sayfasından metin çıkarılamadı.")
    return text


# --------------------------------------------------------------------------- #
# Path / Uploaded file giriş noktaları
# --------------------------------------------------------------------------- #
def _dispatch_by_extension(path: Path) -> str:
    """Uzantıya göre uygun çıkarıcıyı seçer."""
    ext = path.suffix.lower()
    kind = SUPPORTED_EXTENSIONS.get(ext)
    if kind == "pdf":
        return extract_text_from_pdf(path)
    if kind == "docx":
        return extract_text_from_docx(path)
    if kind in ("txt", "md"):
        return extract_text_from_txt(path)
    if kind == "html":
        return extract_text_from_html(path)
    raise DocumentReadError(
        f"Desteklenmeyen dosya tipi: '{ext}'. "
        f"Desteklenenler: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def read_document_from_path(path: Path) -> str:
    """Disk üzerindeki bir dosyadan temizlenmiş metin döndürür."""
    path = Path(path)
    if not path.exists():
        raise DocumentReadError(f"Dosya bulunamadı: {path}")
    logger.info("Dosya okunuyor: %s", path.name)
    return clean_text(_dispatch_by_extension(path))


def read_document_from_uploaded_file(uploaded_file) -> str:
    """Streamlit UploadedFile nesnesinden metin çıkarır.

    Geçici dosya oluşturulur, okunur ve güvenli şekilde silinir.
    """
    name = getattr(uploaded_file, "name", "uploaded")
    ext = Path(name).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError(
            f"Desteklenmeyen dosya tipi: '{ext}'. "
            f"Desteklenenler: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        data = uploaded_file.getvalue()
    except Exception as exc:
        raise DocumentReadError(f"Yüklenen dosya okunamadı: {exc}") from exc

    if not data:
        raise DocumentReadError("Yüklenen dosya boş.")

    tmp_path = save_temp_file(data, ext)
    try:
        return read_document_from_path(tmp_path)
    finally:
        cleanup_temp_file(tmp_path)


# --------------------------------------------------------------------------- #
# URL okuma
# --------------------------------------------------------------------------- #
def _download(url: str) -> requests.Response:
    """URL'yi indirir; hata durumunda DocumentReadError yükseltir."""
    try:
        response = requests.get(
            url,
            timeout=_REQUEST_TIMEOUT,
            headers={"User-Agent": _USER_AGENT},
            stream=True,
        )
        response.raise_for_status()
    except requests.exceptions.MissingSchema as exc:
        raise DocumentReadError(
            "Geçersiz URL. 'http://' veya 'https://' ile başlamalı."
        ) from exc
    except requests.exceptions.Timeout as exc:
        raise DocumentReadError("URL'ye erişim zaman aşımına uğradı.") from exc
    except requests.exceptions.ConnectionError as exc:
        raise DocumentReadError("URL'ye bağlanılamadı (bağlantı hatası).") from exc
    except requests.exceptions.HTTPError as exc:
        code = getattr(exc.response, "status_code", "?")
        raise DocumentReadError(f"URL hata döndürdü (HTTP {code}).") from exc
    except requests.exceptions.RequestException as exc:
        raise DocumentReadError(f"URL indirilemedi: {exc}") from exc
    return response


def _read_response_bytes(response: requests.Response) -> bytes:
    """Response gövdesini boyut sınırıyla okur."""
    content = bytearray()
    for chunk in response.iter_content(chunk_size=65536):
        content.extend(chunk)
        if len(content) > _MAX_DOWNLOAD_BYTES:
            raise DocumentReadError("Dosya çok büyük (60 MB sınırı aşıldı).")
    return bytes(content)


def _looks_like_pdf(url: str, content_type: str, head: bytes) -> bool:
    """İçeriğin PDF olup olmadığını URL, content-type ve imzaya göre tahmin eder."""
    if "application/pdf" in content_type:
        return True
    if urlparse(url).path.lower().endswith(".pdf"):
        return True
    return head[:5] == b"%PDF-"


def _find_pdf_link(html: str, base_url: str) -> str | None:
    """HTML içinden ilk uygun PDF linkini bulur ve mutlak URL'ye çevirir."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    for anchor in soup.find_all("a", href=True):
        href = anchor["href"].strip()
        if href.lower().split("?")[0].endswith(".pdf"):
            return urljoin(base_url, href)
    return None


def read_document_from_url(url: str) -> str:
    """URL'den belge okur.

    - URL doğrudan PDF ise indirip okur.
    - HTML ise metni çıkarır.
    - HTML içinde PDF linki varsa ilk PDF linkini indirip okur.
    """
    url = (url or "").strip()
    if not url:
        raise DocumentReadError("URL boş olamaz.")
    if not url.lower().startswith(("http://", "https://")):
        raise DocumentReadError("URL 'http://' veya 'https://' ile başlamalı.")

    logger.info("URL okunuyor: %s", url)
    response = _download(url)
    content_type = response.headers.get("Content-Type", "").lower()
    data = _read_response_bytes(response)

    # 1) Doğrudan PDF mi?
    if _looks_like_pdf(url, content_type, data):
        return _read_pdf_bytes(data)

    # 2) HTML olarak değerlendir
    try:
        html = data.decode("utf-8", errors="ignore")
    except Exception as exc:  # pragma: no cover
        raise DocumentReadError(f"URL içeriği çözümlenemedi: {exc}") from exc

    # 2a) HTML içinde PDF linki ara
    pdf_link = _find_pdf_link(html, url)
    if pdf_link:
        logger.info("HTML içinde PDF linki bulundu: %s", pdf_link)
        pdf_response = _download(pdf_link)
        pdf_data = _read_response_bytes(pdf_response)
        if _looks_like_pdf(
            pdf_link, pdf_response.headers.get("Content-Type", "").lower(), pdf_data
        ):
            return _read_pdf_bytes(pdf_data)

    # 2b) Düz HTML metni döndür
    return clean_text(_html_to_text(html))


def _read_pdf_bytes(data: bytes) -> str:
    """Bellekteki PDF baytlarını geçici dosya üzerinden okur."""
    tmp_path = save_temp_file(data, ".pdf")
    try:
        return clean_text(extract_text_from_pdf(tmp_path))
    finally:
        cleanup_temp_file(tmp_path)
